import pandas as pd
import numpy as np
import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document  # برای ایجاد فایل Word
from docx.shared import Inches
# 
# 1- خواندن فایل اکسل
file_path = 'Melting_Profile.xlsx'
df = pd.read_excel(file_path)

# 2- حذف قسمت اعشاری برای ستون‌های '[Heat Number]' و '[Power]'
df['[Heat Number]'] = df['[Heat Number]'].apply(lambda x: int(x) if isinstance(x, (int, float)) else x)
df['[Power]'] = df['[Power]'].apply(lambda x: int(x) if isinstance(x, (int, float)) else x)

# 3- اضافه کردن 4000000 به مقادیر ستون '[Heat Number]'
df['[Heat Number]'] = df['[Heat Number]'] + 4000000

# 4- فیلتر کردن داده‌ها بر اساس مقادیر ستون '[Heat Number]'
filtered_data = {}
unique_A_values = df['[Heat Number]'].unique()

for value in unique_A_values[1:-1]:  # حذف شماره ذوب اول و آخر
    filtered_data[value] = df[df['[Heat Number]'] == value].reset_index(drop=True)

# ادامه عملیات برای هر مجموعه داده فیلتر شده
for key, data in filtered_data.items():
    # بررسی ستون '[Power & SC1]' و تغییر مقادیر ستون‌های '[Power]' تا '[SV13]' در صورت صفر بودن
    data.loc[(data['[Power]'] == 0) | (data['[SC1]'] <10) , '[Power]':'[SV31]'] = 0

    # پیدا کردن مقدار ماکزیمم در ستون '[CDRI Total]' و صفر کردن اعداد بعد از آن
    max_index_S = data['[CDRI Total]'].idxmax()
    data.loc[max_index_S + 1:, '[CDRI Total]'] = 0

    # تکرار عملیات برای ستون‌های دیگر
    for col in ['[HDRI Total]', '[Lime Total]', '[Coke Total]', '[DOLOMITE Total]']:
        max_index = data[col].idxmax()
        data.loc[max_index + 1:, col] = 0

    # اضافه کردن ستون جریان برای هر ستون مرتبط
    for flow_col, total_col in [('Total DRI flow rate', '[Total DRI ( HC)]'),
                                ('CDRI Total flow rate', '[CDRI Total]'),
                                ('HDRI Total flow rate', '[HDRI Total]'),
                                ('Lime Total flow rate', '[Lime Total]'),
                                ('Coke Total flow rate', '[Coke Total]'),
                                ('Dolomite Total flow rate', '[DOLOMITE Total]')]:
        data[flow_col] = 0
        for row in range(1, len(data)):
            new_value = (data.iloc[row, data.columns.get_loc(total_col)] - data.iloc[row - 1, data.columns.get_loc(total_col)]) / 10 * 60
            data.iloc[row, data.columns.get_loc(flow_col)] = max(0, new_value)

    # اضافه کردن ستون زمان
    time_interval = 10
    data['Time'] = np.arange(0, len(data) * time_interval, time_interval) / 60

    # اضافه کردن ستون جدید 'Oxygen/ min' بر اساس داده‌های '[KT Oxygen]'
    data['Oxygen/ min'] = data['[KT Oxygen]'] / 60

# طراحی رابط کاربری (GUI) با استفاده از tkinter
def save_heat_number():
    try:
        heat_number = int(combo_box.get())
        if heat_number in filtered_data:
            output_df = filtered_data[heat_number]
            
            # قرار دادن ستون زمان به عنوان اولین ستون
            cols = ['Time'] + [col for col in output_df.columns if col != 'Time']
            output_df = output_df[cols]
            
            # انتخاب مکان برای ذخیره فایل
            save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            
            if save_path:
                output_df.to_excel(save_path, index=False)
                messagebox.showinfo("موفقیت", f"داده‌های مربوط به heat number {heat_number} با موفقیت ذخیره شدند.")
            else:
                messagebox.showwarning("لغو", "ذخیره‌سازی لغو شد.")
        else:
            messagebox.showerror("خطا", "عدد وارد شده وجود ندارد.")
    except ValueError:
        messagebox.showerror("خطا", "لطفاً یک عدد معتبر انتخاب کنید.")

def exit_program():
    root.quit()

# رسم نمودار Melting Profile
def plot_graph_melting():
    heat_number = int(combo_box.get())
    if heat_number in filtered_data:
        data = filtered_data[heat_number]
        
    # محاسبه میانگین متحرک
        data['HDRI Total flow rate (MA)'] = data['HDRI Total flow rate'].rolling(window=5).mean()
        data['CDRI Total flow rate (MA)'] = data['CDRI Total flow rate'].rolling(window=0).mean()
        fig, ax1 = plt.subplots(figsize=(16, 8))
        
        # محور X
        time = data['Time']
        
        # محور Y1 (Power)
        power = data['[Power]']
        ax1.set_xlabel('Time (min)')
        ax1.set_ylabel('Power(MW)', color='tab:red')
        ax1.plot(time, power, label='Power', color='tab:red', linewidth=3)
        ax1.tick_params(axis='y', labelcolor='tab:red')
        ax1.set_ylim([0, 160])
        # محور Y2 (TAP, CDRI Total flow rate, HDRI Total flow rate)
        ax2 = ax1.twinx()
        ax2.set_ylabel('Charging Regime (t/min) & TAP Position', color='tab:blue')
        ax2.plot(time, data['[TAP]'], label='TAP', color='tab:blue', linewidth=2)
        ax2.plot(time, data['CDRI Total flow rate (MA)'], label='CDRI', color='tab:green')
        ax2.plot(time, data['HDRI Total flow rate (MA)'], label='HDRI', color='tab:orange')
        ax2.tick_params(axis='y', labelcolor='tab:blue')
        ax2.set_ylim([0, 15])
        # عنوان نمودار و تنظیم legend
        fig.suptitle(f"Melting Profile for Heat Number {heat_number}")
        fig.legend(loc="upper right", bbox_to_anchor=(1,1), bbox_transform=ax1.transAxes)

        # نمایش نمودار در GUI
        canvas = FigureCanvasTkAgg(fig, master=root)
        canvas.draw()
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
        
        # دکمه برای ذخیره‌سازی نمودار
        save_button = tk.Button(root, text="ذخیره نمودار", command=lambda: save_graph(fig))
        save_button.pack(pady=10)
        
        # دکمه بستن نمودار
        close_button = tk.Button(root, text="بستن نمودار", command=canvas.get_tk_widget().destroy)
        close_button.pack(pady=10)
        
    else:
        messagebox.showerror("خطا", "عدد وارد شده وجود ندارد.")
        
def close_plot(canvas, save_button, close_button):
   canvas.get_tk_widget().destroy()
   save_button.destroy()
   close_button.destroy()

# رسم نمودار Injection Profile
def plot_graph_injection():
    heat_number = int(combo_box.get())
    if heat_number in filtered_data:
        data = filtered_data[heat_number]
        
        fig, ax1 = plt.subplots(figsize=(16,8))
        
        # محور X
        time = data['Time']
        
        # محور Y1 (KT Carbon)
        kt_carbon = data['[KT Carbon]']
        ax1.set_xlabel('Time (min)')
        ax1.set_ylabel('Carbon Injection Flow Rate (kg/min)', color='tab:red')
        ax1.plot(time, kt_carbon, label='KTC', color='tab:red', linewidth=3)
        ax1.tick_params(axis='y', labelcolor='tab:red')
        ax1.set_ylim([0, 70])
        # محور Y2 (Lime, Coke, Dolomite, Oxygen)
        ax2 = ax1.twinx()
        ax2.set_ylabel('Oxygen inj. & Lime, Coke, Dolo charging', color='tab:blue')
        ax2.plot(time, data['Lime Total flow rate'], label='Lime (kg/min)', color='tab:blue', linewidth=3)
        ax2.plot(time, data['Coke Total flow rate'], label='Coke (kg/min)', color='tab:green', linewidth=2)
        ax2.plot(time, data['Dolomite Total flow rate'], label='Dolomite(kg/min)', color='tab:orange', linewidth=2)
        ax2.plot(time, data['Oxygen/ min'], label='KTO (Nm^3/min)', color='tab:purple', linewidth=3)
        ax2.tick_params(axis='y', labelcolor='tab:blue')
        ax2.set_ylim([0, 700])
        # عنوان نمودار و تنظیم legend
        fig.suptitle(f"Injection Profile for Heat Number {heat_number}")
        fig.legend(loc="upper right", bbox_to_anchor=(1,1), bbox_transform=ax1.transAxes)

        # نمایش نمودار در GUI
        canvas = FigureCanvasTkAgg(fig, master=root)
        canvas.draw()
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)
        
        # دکمه برای ذخیره‌سازی نمودار
        save_button = tk.Button(root, text="ذخیره نمودار", command=lambda: save_graph(fig))
        save_button.pack(pady=10)
        
        # دکمه بستن نمودار
        close_button = tk.Button(root, text="بستن نمودار", command=canvas.get_tk_widget().destroy)
        close_button.pack(pady=10)
        
    else:
        messagebox.showerror("خطا", "عدد وارد شده وجود ندارد.")

def close_plot(canvas, save_button, close_button):
    canvas.get_tk_widget().destroy()
    save_button.destroy()
    close_button.destroy()

# تابع ذخیره نمودار
def save_graph(fig):
    save_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png")])
    if save_path:
        fig.savefig(save_path)

# تابع ذخیره همه نمودارها در یک فایل Word
def save_all_graphs_to_word():
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if save_path:
        doc = Document()
        for heat_number in filtered_data:
            data = filtered_data[heat_number]
# نمودار Melting Profile
            data['HDRI Total flow rate (MA)'] = data['HDRI Total flow rate'].rolling(window=5).mean()
            data['CDRI Total flow rate (MA)'] = data['CDRI Total flow rate'].rolling(window=0).mean()
            fig, ax1 = plt.subplots(figsize=(16, 8))
            time = data['Time']
            power = data['[Power]']
            ax1.set_xlabel('Time (min)')
            ax1.set_ylabel('Power(MW)', color='tab:red')
            ax1.plot(time, power, label='Power', color='tab:red', linewidth=3)
            ax1.tick_params(axis='y', labelcolor='tab:red')
            ax1.set_ylim([0, 160])
            ax2 = ax1.twinx()
            ax2.set_ylabel('Charging Regime (t/min) & TAP Position', color='tab:blue')
            ax2.plot(time, data['[TAP]'], label='TAP', color='tab:blue', linewidth=2)
            ax2.plot(time, data['CDRI Total flow rate (MA)'], label='CDRI', color='tab:green')
            ax2.plot(time, data['HDRI Total flow rate (MA)'], label='HDRI', color='tab:orange')
            ax2.tick_params(axis='y', labelcolor='tab:blue')
            ax2.set_ylim([0, 15])
            fig.suptitle(f"Melting Profile for Heat Number {heat_number}")
            fig.legend(loc="upper right", bbox_to_anchor=(1,1), bbox_transform=ax1.transAxes)
            fig.savefig('melting_temp.png')
            doc.add_heading(f'Melting Profile for Heat Number {heat_number}', level=1)
            doc.add_picture('melting_temp.png', width=Inches(6))
            plt.close(fig)
            
            # نمودار Injection Profile
            fig, ax1 = plt.subplots(figsize=(16, 8))
            kt_carbon = data['[KT Carbon]']
            ax1.set_xlabel('Time (min)')
            ax1.set_ylabel('Carbon Injection Flow Rate (kg/min)', color='tab:red')
            ax1.plot(time, kt_carbon, label='KTC', color='tab:red', linewidth=3)
            ax1.tick_params(axis='y', labelcolor='tab:red')
            ax1.set_ylim([0, 70])
            ax2 = ax1.twinx()
            ax2.set_ylabel('Oxygen inj. & Lime, Coke, Dolo charging', color='tab:blue')
            ax2.plot(time, data['Lime Total flow rate'], label='Lime (kg/min)', color='tab:blue', linewidth=3)
            ax2.plot(time, data['Coke Total flow rate'], label='Coke (kg/min)', color='tab:green', linewidth=2)
            ax2.plot(time, data['Dolomite Total flow rate'], label='Dolomite(kg/min)', color='tab:orange', linewidth=2)
            ax2.plot(time, data['Oxygen/ min'], label='KTO (Nm^3/min)', color='tab:purple', linewidth=3)
            ax2.tick_params(axis='y', labelcolor='tab:blue')
            ax2.set_ylim([0, 700])
            fig.suptitle(f"Injection Profile for Heat Number {heat_number}")
            fig.legend(loc="upper right", bbox_to_anchor=(1,1), bbox_transform=ax1.transAxes)
            fig.savefig('injection_temp.png')
            doc.add_heading(f'Injection Profile for Heat Number {heat_number}', level=1)
            doc.add_picture('injection_temp.png', width=Inches(6))
            plt.close(fig)

        doc.save(save_path)
        messagebox.showinfo("موفقیت", "همه نمودارها با موفقیت در فایل Word ذخیره شدند.")

# رابط کاربری اصلی
root = tk.Tk()
root.title("Melting & Injection Profile Plotter")

# ایجاد منوی کشویی برای شماره‌های ذوب
heat_numbers = [str(num) for num in unique_A_values[1:-1]]  # حذف اولین و آخرین شماره ذوب
combo_box = ttk.Combobox(root, values=heat_numbers, state="readonly")
combo_box.pack(pady=5)

# دکمه‌های رسم نمودار و ذخیره
plot_button_melting = tk.Button(root, text="رسم نمودار Melting Profile", command=plot_graph_melting)
plot_button_melting.pack(pady=10)

plot_button_injection = tk.Button(root, text="رسم نمودار Injection Profile", command=plot_graph_injection)
plot_button_injection.pack(pady=10)

save_button = tk.Button(root, text="ذخیره داده‌ها", command=save_heat_number)
save_button.pack(pady=10)

save_all_button = tk.Button(root, text="ذخیره همه نمودارها در فایل Word", command=save_all_graphs_to_word)
save_all_button.pack(pady=10)

# خروج از برنامه
exit_button = tk.Button(root, text="خروج", command=exit_program)
exit_button.pack(pady=10)

root.mainloop()

